import io

from app.models.enums import SourceType
from app.parsing.base import ParsedDocument, TextBlock, register


@register(SourceType.docx)
class DocxParser:
    def parse(self, *, data: bytes | None, url: str | None, filename: str) -> ParsedDocument:
        import docx  # python-docx

        if data is None:
            raise ValueError("DOCX parser requires file bytes")
        document = docx.Document(io.BytesIO(data))
        parts = [p.text for p in document.paragraphs if p.text.strip()]
        # Include table cell text as well.
        for table in document.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells if c.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))
        text = "\n".join(parts)
        return ParsedDocument(blocks=[TextBlock(text=text)], meta={"parser": "python-docx"})
